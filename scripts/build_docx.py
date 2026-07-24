"""manual-draft.md → 매뉴얼 DOCX 생성기 (전용 docx skill 이 없는 환경의 표준 폴백).

output-formats.md 의 docx 명세를 내장한다:
  표지(분리 세트는 --cover-label 2행 규격) → 자동 목차(TOC 필드) → Heading 1(장)/
  2(절)/3(하위 절), 각 장 새 페이지, 본문 맑은 고딕, 이미지 폭 맞춤(세로로 긴
  캡처는 높이 상한에서 축소) + 검은 실선 테두리 + [사진 N] 캡션, placeholder 회색
  안내 문단, 머리글(문서명)·바닥글(가운데 페이지 번호), 마크다운 서식은 실제 서식으로 변환.

사용 예:
  python build_docx.py --draft manual-work/manual-draft.md \
      --screenshots manual-work/screenshots --out 사용자매뉴얼_시스템명_20260711.docx

주의: 자동 목차는 Word 에서 열어 F9(필드 갱신) 시 채워진다.
종료 코드: 0 성공 / 1 파싱·생성 오류 / 2 python-docx 미설치
"""

import argparse
import os
import re
import sys

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from draft_parser import (parse_draft, parse_inline, parse_meta, image_size,
                          resolve_image, tile_tall_image)


def fail(msg, code=1):
    print(f"[build_docx] 오류: {msg}", file=sys.stderr)
    sys.exit(code)


try:
    import docx
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    fail("python-docx 가 설치되어 있지 않습니다. 설치 후 재시도하세요:\n  pip install python-docx", code=2)

FONT = "맑은 고딕"
NOTE = RGBColor(0xC0, 0x39, 0x2B)
MUTED = RGBColor(0x8A, 0x8F, 0x98)
PH = RGBColor(0x6B, 0x72, 0x80)


def set_kfont(run, size=None, bold=None, color=None):
    run.font.name = FONT
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_runs(para, segments, size=10.5, color=None, bold=False):
    if isinstance(segments, str):
        segments = [(segments, bold)]
    for text, seg_bold in segments:
        r = para.add_run(text)
        set_kfont(r, size=size, bold=bold or seg_bold, color=color)
    return para


def add_field(para, instr):
    """Word 필드(TOC, PAGE 등)를 문단에 삽입한다."""
    r = para.add_run()
    for tag, attrs, text in (
        ("w:fldChar", {"w:fldCharType": "begin"}, None),
        ("w:instrText", {"xml:space": "preserve"}, instr),
        ("w:fldChar", {"w:fldCharType": "separate"}, None),
        ("w:fldChar", {"w:fldCharType": "end"}, None),
    ):
        el = OxmlElement(tag)
        for k, v in attrs.items():
            el.set(qn(k), v)
        if text:
            el.text = text
        r._element.append(el)


def shade(para, hex_fill="F1F2F4"):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    pPr.append(shd)


IMG_W_CM = 16.0       # 스크린샷 삽입 폭
IMG_MAX_H_CM = 20.0   # 삽입 높이 상한 — A4 본문 가용(25.7cm) 안에서 잘리지 않게


def apply_picture_border(inline_shape):
    """그림 서식: 검은색 실선 테두리(0.75pt) — pptx 와 동일한 이미지 규격.
    흰 배경 캡처가 흰 용지에 경계 없이 녹아드는 것을 막는다."""
    spPr = inline_shape._inline.graphic.graphicData.pic.spPr
    ln = OxmlElement("a:ln")
    ln.set("w", "9525")  # 0.75pt (EMU)
    fill = OxmlElement("a:solidFill")
    clr = OxmlElement("a:srgbClr")
    clr.set("val", "000000")
    fill.append(clr)
    ln.append(fill)
    spPr.append(ln)


def add_screenshot(para, path):
    """스크린샷 삽입 — 폭 16cm 기준, 세로로 긴 캡처는 높이 상한에서 비율 유지 축소."""
    size = image_size(path)
    run = para.add_run()
    if size and size[0] and IMG_W_CM * size[1] / size[0] > IMG_MAX_H_CM:
        shp = run.add_picture(path, height=Cm(IMG_MAX_H_CM))
    else:
        shp = run.add_picture(path, width=Cm(IMG_W_CM))
    apply_picture_border(shp)


def render_blocks(doc_x, blocks, draft_dir, shots_dir, counters):
    for b in blocks:
        t = b["type"]
        if t == "para":
            add_runs(doc_x.add_paragraph(), parse_inline(b["text"]))
        elif t == "access":
            add_runs(doc_x.add_paragraph(), [("접근 경로: ", True), (b["text"], False)], size=10.5)
        elif t == "note":
            add_runs(doc_x.add_paragraph(), parse_inline(b["text"]), color=NOTE, bold=True)
        elif t == "bullets":
            for it in b["items"]:
                add_runs(doc_x.add_paragraph(style="List Bullet"), parse_inline(it))
        elif t == "numbered":
            # 파서가 보존한 원본 마커를 그대로 렌더한다 — 블록이 분절돼도 재번호하지
            # 않아 원고(=배지) 번호와 항상 일치한다
            for it in b["items"]:
                add_runs(doc_x.add_paragraph(), [(f"{it['marker']} ", False)] + parse_inline(it["text"]))
        elif t == "table":
            rows = b["rows"]
            table = doc_x.add_table(rows=len(rows), cols=max(len(r) for r in rows))
            table.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    cell = table.cell(ri, ci)
                    cell.text = ""
                    add_runs(cell.paragraphs[0], cell_text, size=10, bold=(ri == 0))
                    if ri == 0:
                        tcPr = cell._tc.get_or_add_tcPr()
                        shd = OxmlElement("w:shd")
                        shd.set(qn("w:val"), "clear")
                        shd.set(qn("w:fill"), "E8ECF4")
                        tcPr.append(shd)
            doc_x.add_paragraph()
        elif t == "image":
            path = resolve_image(b["src"], draft_dir, shots_dir)
            # 세로로 긴 캡처는 폭을 줄이는 대신 표준 비율 밴드로 타일 분할한다 —
            # 폭 16cm가 균일성 앵커이므로 모든 밴드가 동일 폭으로 들어간다(요소 경계 절단).
            targets = []
            if path:
                size = image_size(path)
                if size and size[0] and size[1] / size[0] > IMG_MAX_H_CM / IMG_W_CM:
                    targets = tile_tall_image(path, shots_dir)
                if not targets:
                    targets = [path]
            if not targets:
                p = doc_x.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_runs(p, f"(이미지 파일 누락: {b['src']})", color=NOTE)
            for ti, tp in enumerate(targets):
                p = doc_x.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_screenshot(p, tp)
            if b.get("caption"):
                suffix = f" ({len(targets)}분할)" if len(targets) > 1 else ""
                cap = doc_x.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_runs(cap, b["caption"] + suffix, size=9, color=MUTED)
        elif t == "placeholder":
            p = doc_x.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shade(p)
            add_runs(p, f"화면 이미지 추후 삽입 — {b['scr']} {b['name']}", size=10.5, color=PH, bold=True)
            cap = doc_x.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs(cap, f"[사진 -] {b['name']} (추후 삽입)", size=9, color=MUTED)


def main():
    ap = argparse.ArgumentParser(description="manual-draft.md → 매뉴얼 DOCX")
    ap.add_argument("--draft", required=True)
    ap.add_argument("--screenshots", help="스크린샷 디렉토리 (기본: draft 위치의 screenshots/)")
    ap.add_argument("--out", required=True)
    ap.add_argument("--title", help="표지 제목 (기본: 원고 # 제목)")
    ap.add_argument("--audience")
    ap.add_argument("--version")
    ap.add_argument("--date")
    ap.add_argument("--cover-label",
                    help='표지 1행 매뉴얼 구분 표기(예: "관리자 매뉴얼") — 독자 영역별로 '
                         "매뉴얼이 분리된 세트일 때만 지정. 지정 시 표지가 [구분]+[프로젝트명] "
                         "2행이 되고, 미지정 시 프로젝트명만 표기한다")
    ap.add_argument("--skip-validate", action="store_true", help="원고 사전 검증을 건너뛴다")
    args = ap.parse_args()

    if not os.path.exists(args.draft):
        fail(f"원고 없음: {args.draft}")
    draft_dir = os.path.dirname(os.path.abspath(args.draft))
    shots_dir = args.screenshots or os.path.join(draft_dir, "screenshots")

    doc = parse_draft(args.draft)
    if not doc["chapters"]:
        fail("장(## NN. 제목)을 찾지 못했습니다 — manual-template.md 규약을 확인하세요")

    # 원고 사전 검증 게이트 (build_pptx 와 동일)
    if not args.skip_validate:
        import validate_draft
        with open(args.draft, encoding="utf-8-sig") as f:
            raw = f.read()
        errors, warns = validate_draft.validate(doc, draft_dir, shots_dir, raw_text=raw)
        for w in warns:
            print(f"[build_docx] 원고 WARN: {w}")
        if errors:
            for e in errors:
                print(f"[build_docx] 원고 ERROR: {e}", file=sys.stderr)
            fail(f"원고 검증 실패({len(errors)}건) — 원고를 수정하거나 --skip-validate 로 우회하세요")

    audience, version, date = parse_meta(doc["meta"])
    audience = args.audience or audience
    version = args.version or version
    date = args.date or date
    title = args.title or doc["title"] or "사용자 매뉴얼"

    doc_x = docx.Document()
    style = doc_x.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for name, size in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
        h = doc_x.styles[name]
        h.font.name = FONT
        h.font.size = Pt(size)
        h.element.get_or_add_rPr()
        rf = h.element.rPr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            h.element.rPr.append(rf)
        rf.set(qn("w:eastAsia"), FONT)

    section = doc_x.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2.0)
    section.left_margin = section.right_margin = Cm(2.5)
    section.different_first_page_header_footer = True  # 표지에는 머리글·페이지 번호 없음
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_runs(hp, f"{title}", size=8, color=MUTED)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_kfont(fp.add_run(""), size=9)
    add_field(fp, "PAGE")

    # 표지 — pptx 와 동일 규격: 분리 세트는 [매뉴얼 구분]+[프로젝트명] 2행,
    # 단일 매뉴얼은 프로젝트명 + 대상 부제
    cover_title = title
    m = re.match(r"^(.{0,20}매뉴얼)\s*[—\-–:]\s*(.+)$", cover_title)
    if m:
        cover_title = m.group(2).strip()
    for _ in range(6):
        doc_x.add_paragraph()
    if args.cover_label:
        p = doc_x.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs(p, args.cover_label, size=16, color=MUTED)
    p = doc_x.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(p, cover_title, size=28, bold=True)
    if not args.cover_label:
        p = doc_x.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs(p, f"{audience} 사용자 매뉴얼" if "매뉴얼" not in audience else audience,
                 size=13, color=MUTED)
    p = doc_x.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(p, " · ".join(v for v in (audience, f"버전 {version}" if version else "", date) if v), size=12, color=MUTED)
    doc_x.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # 자동 목차
    p = doc_x.add_paragraph()
    add_runs(p, "목차", size=16, bold=True)
    add_field(doc_x.add_paragraph(), r'TOC \o "1-3" \h \z \u')
    tip = doc_x.add_paragraph()
    add_runs(tip, "※ 목차가 비어 보이면 Word에서 Ctrl+A 후 F9로 필드를 갱신하세요.", size=8, color=MUTED)
    doc_x.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    counters = {}
    for ci, ch in enumerate(doc["chapters"]):
        if ci > 0:
            doc_x.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        h = doc_x.add_heading(f"{ch['num']}. {ch['title']}", level=1)
        for r in h.runs:
            set_kfont(r)
        render_blocks(doc_x, ch["intro"], draft_dir, shots_dir, counters)
        for sec in ch["sections"]:
            level = 3 if sec["num"].count(".") >= 2 else 2
            h = doc_x.add_heading(f"{sec['num']} {sec['title']}", level=level)
            for r in h.runs:
                set_kfont(r)
            render_blocks(doc_x, sec["blocks"], draft_dir, shots_dir, counters)

    os.makedirs(os.path.dirname(os.path.abspath(args.out)), exist_ok=True)
    doc_x.save(args.out)
    print(f"[build_docx] 저장 완료: {args.out}")


if __name__ == "__main__":
    main()
